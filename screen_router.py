"""
routers/screen_router.py — Resume Screening Routes
=====================================================
Uses Google Gemini API — completely free tier, no dollar payment needed.
Free limit: 1,500 requests/day which is plenty for a growing business.
"""

import base64
import io
import json
import logging
import os
import traceback
from pathlib import Path

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile, Request
from fastapi.responses import JSONResponse
from sqlalchemy.orm import Session

import security as auth
import requests
import models
from database import get_db

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api", tags=["screening"])


# ── Text Extraction ───────────────────────────────────────────────────────────

def extract_text(filename: str, file_bytes: bytes) -> str:
    """Extract plain text from PDF, DOCX, or TXT files."""
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages = [p.extract_text() for p in pdf.pages if p.extract_text()]
        return "\n\n".join(pages)

    elif ext == ".docx":
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        texts.append(cell.text.strip())
        return "\n".join(texts)

    elif ext == ".txt":
        return file_bytes.decode("utf-8", errors="ignore")

    raise ValueError(
        f"Unsupported file type: {ext}. Please upload PDF, DOCX, or TXT.")


# ── AI Screening Prompt ───────────────────────────────────────────────────────

PROMPT = """You are a world-class talent acquisition specialist with 20+ years of experience.

Analyse this resume against the job description and return ONLY valid JSON — no markdown, no extra text.

JOB DESCRIPTION:
{job_description}

RESUME (Candidate: {candidate_name}):
{resume_text}

Return ONLY this JSON structure with no other text before or after it:
{{
  "candidate_name": "Full name from resume or '{candidate_name}'",
  "overall_score": <integer 0-100>,
  "recommendation": "Strong Hire",
  "executive_summary": "<2-3 sentence summary of candidate fit>",
  "scores": {{
    "technical_skills": <0-100>,
    "experience": <0-100>,
    "education": <0-100>,
    "cultural_fit": <0-100>,
    "leadership": <0-100>,
    "communication": <0-100>
  }},
  "strengths": ["<strength 1>", "<strength 2>", "<strength 3>", "<strength 4>"],
  "gaps": ["<gap 1>", "<gap 2>", "<gap 3>"],
  "key_skills": ["<skill1>", "<skill2>", "<skill3>", "<skill4>", "<skill5>"],
  "experience_years": <integer>,
  "highest_education": "<degree and field>",
  "standout_achievements": "<most impressive achievement>",
  "interview_questions": ["<question 1>", "<question 2>", "<question 3>"],
  "red_flags": [],
  "salary_expectation": "<estimated range>",
  "availability_signals": "<notice period or availability>"
}}

The recommendation field must be one of: "Strong Hire", "Hire", "Maybe", "No Hire"."""


def screen_with_openai(api_key: str, job_description: str,
                       resume_text: str, candidate_name: str) -> dict:
    """Use OpenAI chat completions (Deepseek/OpenAI) to screen a resume."""
    try:
        from openai import OpenAI
    except ModuleNotFoundError as e:
        raise RuntimeError(
            "OpenAI SDK is not installed. Install it with: pip install openai"
        ) from e

    # Allow using Deepseek, Groq, or other OpenAI-compatible providers by setting a custom API base.
    api_base = os.getenv("OPENAI_API_BASE") or os.getenv("DEEPSEEK_API_BASE")
    groq_base = os.getenv("GROQ_API_BASE")
    groq_key = os.getenv("GROQ_API_KEY")

    # If Groq credentials are present, prefer direct REST call (their Postman example).
    if groq_base and groq_key:
        # Normalize base to drop any trailing '/models' or '/'
        base = groq_base.rstrip('/')
        if base.endswith('/models'):
            base = base.rsplit('/models', 1)[0]
        endpoint = f"{base}/chat/completions"
        headers = {
            "Authorization": f"Bearer {groq_key}",
            "Content-Type": "application/json",
        }
        model = os.getenv("GROQ_MODEL_NAME") or os.getenv(
            "DEEPSEEK_MODEL_NAME") or os.getenv("OPENAI_MODEL_NAME") or "gpt-4o-mini"
        payload = {
            "model": model,
            "messages": [
                {"role": "system",
                    "content": "You are a world-class talent acquisition specialist."},
                {"role": "user", "content": PROMPT.format(
                    job_description=job_description, resume_text=resume_text[:8000], candidate_name=candidate_name)},
            ],
            "temperature": 0.2,
        }
        resp = requests.post(endpoint, headers=headers,
                             json=payload, timeout=60)
        if resp.status_code == 200:
            try:
                body = resp.json()
                content = body["choices"][0]["message"]["content"]
                raw = (content or "").strip()
                if raw.startswith("```"):
                    lines = raw.split("\n")
                    raw = "\n".join(
                        lines[1:-1] if lines[-1].strip() == "```" else lines[1:])
                return json.loads(raw)
            except Exception as e:
                raise RuntimeError(
                    f"Failed to parse Groq response: {e}\nBody: {resp.text}") from e
        else:
            # Map auth error to clearer message
            if resp.status_code == 401:
                raise RuntimeError(
                    "Authentication with Groq failed (401). Check GROQ_API_KEY and GROQ_API_BASE.")
            raise RuntimeError(
                f"Groq request failed: {resp.status_code} - {resp.text}")

    # fall back to OpenAI/Deepseek via SDK
    try:
        client = OpenAI(api_key=api_key, base_url=api_base) if api_base else OpenAI(
            api_key=api_key)
    except TypeError:
        # Older openai client versions may not accept api_base in constructor; fall back to env var
        if api_base:
            os.environ["OPENAI_API_BASE"] = api_base
        client = OpenAI(api_key=api_key)

    prompt = PROMPT.format(
        job_description=job_description,
        resume_text=resume_text[:8000],
        candidate_name=candidate_name
    )

    model = os.getenv("OPENAI_MODEL_NAME") or os.getenv("DEEPSEEK_MODEL_NAME") or os.getenv(
        "GROQ_MODEL_NAME") or os.getenv("DEEPSEEK_MODEL_NAME") or "gpt-4o-mini"

    # Build chat messages — keep system short and user prompt detailed
    messages = [
        {"role": "system", "content": "You are a world-class talent acquisition specialist."},
        {"role": "user", "content": prompt},
    ]

    try:
        response = client.chat.completions.create(
            model=model,
            messages=messages,
            temperature=0.2,
        )
    except Exception as e:
        # Improve error clarity for invalid API key / 401 responses from provider
        err_text = str(e)
        if "invalid_api_key" in err_text or "Incorrect API key" in err_text or "401" in err_text:
            raise RuntimeError(
                "Authentication with the AI provider failed (401).\n"
                "If you're using Deepseek, ensure you set `DEEPSEEK_API_BASE` and `DEEPSEEK_API_KEY`.\n"
                "If you're using Groq, set `GROQ_API_BASE` and `GROQ_API_KEY`.\n"
                "If you're using OpenAI, set `OPENAI_API_KEY`.\n"
                "Do NOT share keys publicly — rotate the exposed key if it was committed."
            ) from e
        raise

    # Extract text content
    try:
        choice = response.choices[0]
        # new-style object may have .message or .message['content']
        content = None
        if hasattr(choice, "message"):
            msg = choice.message
            content = msg.get("content") if isinstance(
                msg, dict) else getattr(msg, "content", None)
        else:
            # fallback to text property
            content = getattr(choice, "text", None) or str(choice)
        raw = (content or "").strip()
    except Exception:
        raw = str(response)

    # Strip markdown fences if present
    if raw.startswith("```"):
        lines = raw.split("\n")
        raw = "\n".join(lines[1:-1] if lines[-1].strip()
                        == "```" else lines[1:])

    return json.loads(raw.strip())


# ── Main Screening Route ──────────────────────────────────────────────────────

@router.post("/screen")
async def screen_resumes(
    request: Request,
    job_description: str = Form(...),
    job_title: str = Form(default="Open Position"),
    company_name: str = Form(default=""),
    csrf_token: str = Form(...),
    files: list[UploadFile] = File(...),
    db: Session = Depends(get_db),
    current_user: models.User = Depends(auth.get_current_user),
):
    import main as _main
    _main.require_csrf(request, csrf_token)

    MAX_FILE_SIZE = 10 * 1024 * 1024
    MAX_FILES = 50
    ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}

    if not job_description.strip():
        raise HTTPException(400, "Job description is required.")
    if len(job_description) > 20000:
        raise HTTPException(400, "Job description is too long. Keep it under 20,000 characters.")
    if not files or all(f.filename == "" for f in files):
        raise HTTPException(400, "At least one resume file is required.")
    if len(files) > MAX_FILES:
        raise HTTPException(400, f"Maximum {MAX_FILES} files allowed per screening.")

    for upload in files:
        if not upload.filename:
            continue
        ext = Path(upload.filename).suffix.lower()
        if ext not in ALLOWED_EXTENSIONS:
            raise HTTPException(
                400,
                f"File '{upload.filename}' has unsupported format. Allowed: PDF, DOCX, TXT",
            )
        if upload.size and upload.size > MAX_FILE_SIZE:
            size_mb = upload.size / (1024 * 1024)
            raise HTTPException(
                400,
                f"File '{upload.filename}' is too large ({size_mb:.1f}MB). Maximum: 10MB",
            )

    from datetime import datetime, timedelta

    if current_user.usage_reset_date is None or current_user.usage_reset_date <= datetime.utcnow():
        current_user.screenings_used_this_month = 0
        current_user.usage_reset_date = datetime.utcnow() + timedelta(days=30)
        db.commit()

    FREE_LIMIT = 5
    PRO_LIMIT = 100
    ENTERPRISE_LIMIT = None

    now = datetime.utcnow()
    tier = (current_user.tier or "FREE").upper()
    is_pro = tier == "PRO" and (
        current_user.premium_until is None or current_user.premium_until > now
    )
    is_enterprise = tier == "ENTERPRISE" and (
        current_user.premium_until is None or current_user.premium_until > now
    )

    if is_enterprise:
        monthly_limit = ENTERPRISE_LIMIT
    elif is_pro:
        monthly_limit = PRO_LIMIT
    else:
        monthly_limit = FREE_LIMIT

    if monthly_limit is not None and current_user.screenings_used_this_month >= monthly_limit:
        tier_name = "Enterprise" if is_enterprise else "Pro" if is_pro else "Free"
        raise HTTPException(
            403,
            f"You've reached your {monthly_limit} screenings/month limit for {tier_name} tier. "
            f"Upgrade by bank transfer from your dashboard > Plans.",
        )

    resolved_key = os.environ.get("OPENAI_API_KEY", "").strip() \
        or os.environ.get("DEEPSEEK_API_KEY", "").strip() \
        or os.environ.get("GROQ_API_KEY", "").strip() \
        or os.environ.get("GEMINI_API_KEY", "").strip()
    if not resolved_key:
        raise HTTPException(500, "AI service is not configured. Please contact support.")

    job = db.query(models.Job).filter(
        models.Job.user_id == current_user.id,
        models.Job.title == job_title,
        models.Job.company == company_name,
    ).first()

    if not job:
        job = models.Job(
            user_id=current_user.id,
            title=job_title,
            company=company_name,
            description=job_description,
        )
        db.add(job)
        db.commit()
        db.refresh(job)

    valid_uploads = [upload for upload in files if upload.filename]
    screening = models.Screening(
        user_id=current_user.id,
        job_id=job.id,
        total_files=len(valid_uploads),
        processed_candidates=0,
        total_candidates=0,
        status="QUEUED",
    )
    db.add(screening)
    db.commit()
    db.refresh(screening)

    queued_files, rejected_files = [], []

    for upload in valid_uploads:
        try:
            file_bytes = await upload.read()
            if len(file_bytes) > MAX_FILE_SIZE:
                size_mb = len(file_bytes) / (1024 * 1024)
                rejected_files.append({
                    "file": upload.filename,
                    "error": f"File is too large ({size_mb:.1f}MB). Maximum: 10MB",
                })
                continue

            candidate_name = (
                Path(upload.filename).stem
                .replace("_", " ")
                .replace("-", " ")
                .title()
            )
            candidate_row = models.Candidate(
                screening_id=screening.id,
                filename=upload.filename,
                candidate_name=candidate_name,
                status="QUEUED",
                file_content_b64=base64.b64encode(file_bytes).decode("ascii"),
            )
            db.add(candidate_row)
            db.commit()
            db.refresh(candidate_row)

            queued_files.append({
                "candidate_id": candidate_row.id,
                "filename": upload.filename,
            })
        except Exception as e:
            logger.error("Error queueing %s: %s", upload.filename, traceback.format_exc())
            rejected_files.append({
                "file": upload.filename,
                "error": f"Queueing failed: {str(e)}",
            })

    if not queued_files:
        screening.status = "FAILED"
        screening.error_message = "No resume files could be queued."
        db.commit()
        return JSONResponse(
            status_code=400,
            content={
                "detail": "No resume files could be queued.",
                "screening_id": screening.id,
                "errors": rejected_files,
            },
        )

    screening.total_files = len(queued_files)
    current_user.screenings_used_this_month += 1
    db.commit()

    from screening_tasks import process_screening_task
    task = process_screening_task.delay(screening.id, queued_files)

    return JSONResponse(status_code=202, content={
        "job_title": job_title,
        "company_name": company_name,
        "screening_id": screening.id,
        "task_id": task.id,
        "status": screening.status,
        "total_files": screening.total_files,
        "processed_candidates": screening.processed_candidates,
        "total_processed": 0,
        "total_errors": len(rejected_files),
        "results": [],
        "errors": rejected_files,
    })
# ── History Routes ────────────────────────────────────────────────────────────

@router.get("/history")
async def get_history(
    db: Session = Depends(get_db),
    current_user: models.User = Depends(auth.get_current_user)
):
    screenings = (
        db.query(models.Screening)
        .filter(models.Screening.user_id == current_user.id)
        .order_by(models.Screening.created_at.desc())
        .limit(50)
        .all()
    )
    data = []
    for s in screenings:
        top = max(s.candidates, key=lambda c: c.overall_score or 0, default=None)
        data.append({
            "id": s.id,
            "job_title": s.job.title,
            "company": s.job.company,
            "total_candidates": s.total_candidates,
            "created_at": s.created_at.isoformat(),
            "top_candidate": top.candidate_name if top else None,
            "top_score": top.overall_score if top else None,
        })
    return JSONResponse({"screenings": data})


@router.get("/screening/{screening_id}")
async def get_screening(
    screening_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(auth.get_current_user)
):
    screening = db.query(models.Screening).filter(
        models.Screening.id == screening_id,
        models.Screening.user_id == current_user.id
    ).first()

    if not screening:
        raise HTTPException(404, "Screening not found.")

    candidates = sorted(
        [c.result_json for c in screening.candidates if c.result_json],
        key=lambda x: x.get("overall_score", 0),
        reverse=True
    )
    errors = [
        {"file": c.filename, "error": c.error_message or "Processing failed."}
        for c in screening.candidates
        if c.status == "FAILED"
    ]

    return JSONResponse({
        "job_title": screening.job.title,
        "company_name": screening.job.company or "",
        "screening_id": screening.id,
        "created_at": screening.created_at.isoformat(),
        "status": screening.status,
        "error_message": screening.error_message,
        "total_files": screening.total_files or len(screening.candidates),
        "processed_candidates": screening.processed_candidates or len(candidates) + len(errors),
        "total_processed": len(candidates),
        "total_errors": len(errors),
        "results": candidates,
        "errors": errors
    })
